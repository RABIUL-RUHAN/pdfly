"""
FR-06: Convert PDF to Word (.docx)

Features:
- Preserves layout using pdf2docx
- Extracts tables using pdfplumber
- Supports scanned PDFs using OCR fallback
"""

import os

from pdf2docx import Converter
import pdfplumber

from docx import Document
from docx.shared import Inches

import fitz  # PyMuPDF
from PIL import Image
import pytesseract

from .validators import ValidationError


MIN_TEXT_THRESHOLD = 20


def pdf_has_text(pdf_path):
    """
    Check whether PDF contains extractable text
    """

    text_length = 0

    with pdfplumber.open(pdf_path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            text_length += len(text.strip())

    return text_length >= MIN_TEXT_THRESHOLD



def convert_with_pdf2docx(input_path, output_path):
    """
    Primary conversion method
    """

    try:
        converter = Converter(input_path)

        converter.convert(
            output_path,
            start=0,
            end=None
        )

        converter.close()

        return True

    except Exception:
        return False



def extract_tables(input_path, output_path):
    """
    Extract tables and recreate them in Word
    """

    doc = Document()

    found_table = False


    with pdfplumber.open(input_path) as pdf:

        for page_number, page in enumerate(pdf, start=1):

            tables = page.extract_tables()


            for table in tables:

                if not table:
                    continue


                found_table = True


                doc.add_paragraph(
                    f"Table from page {page_number}"
                )


                rows = len(table)
                cols = len(table[0])


                word_table = doc.add_table(
                    rows=rows,
                    cols=cols
                )


                for r in range(rows):

                    for c in range(cols):

                        value = table[r][c]

                        if value is None:
                            value = ""

                        word_table.cell(
                            r,
                            c
                        ).text = value



    if found_table:

        doc.save(output_path)

        return True


    return False



def convert_scanned_pdf(input_path, output_path):
    """
    OCR conversion for scanned PDFs
    """

    doc = Document()

    pdf = fitz.open(input_path)


    for page_number in range(len(pdf)):

        page = pdf[page_number]


        pix = page.get_pixmap(
            dpi=300
        )


        image = Image.frombytes(
            "RGB",
            [
                pix.width,
                pix.height
            ],
            pix.samples
        )


        text = pytesseract.image_to_string(
            image
        )


        doc.add_paragraph(text)



    doc.save(output_path)

    return True



def pdf_to_word(input_path, output_path):

    if not os.path.exists(input_path):
        raise ValidationError(
            "PDF file does not exist."
        )


    warning = None


    # Method 1: normal PDF conversion

    success = convert_with_pdf2docx(
        input_path,
        output_path
    )


    if success:

        if not pdf_has_text(input_path):

            warning = (
                "PDF appears to be scanned. "
                "OCR may be required."
            )

        return output_path, warning



    # Method 2: table reconstruction

    success = extract_tables(
        input_path,
        output_path
    )


    if success:

        return output_path, (
            "Tables were reconstructed manually."
        )


    # Method 3: OCR fallback

    success = convert_scanned_pdf(
        input_path,
        output_path
    )


    if success:

        return output_path, (
            "OCR conversion used."
        )


    raise ValidationError(
        "Unable to convert this PDF."
    )